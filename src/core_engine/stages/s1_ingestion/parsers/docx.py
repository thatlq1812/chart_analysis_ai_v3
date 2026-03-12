"""
DOCX Document Parser

Handles Microsoft Word (.docx) files via python-docx.

Context extraction strategy:
  - Iterates the document body paragraph-by-paragraph.
  - When a paragraph contains an InlineShape (embedded image),
    collects the preceding N and following N text paragraphs as
    surrounding_text, and checks for nearby caption-like paragraphs.
  - Each embedded image is exported as a PNG byte-stream and converted
    to a BGR numpy array.
  - If the DOCX contains no embedded images, every page placeholder
    is rendered at the given DPI using LibreOffice (if available) or
    the document text body is returned as a single ParsedPage.

Limitations:
  - Floating images (positioned outside paragraph flow) may be missed.
    This is a known limitation of the python-docx XML model.
  - Rendering full DOCX pages at exact DPI requires LibreOffice.
    When LibreOffice is unavailable, each embedded image is returned
    as a standalone ParsedPage (layout-free mode).
"""

from __future__ import annotations

import io
import logging
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple

import cv2
import numpy as np

from .base import BaseDocumentParser, ParsedDocument, ParsedPage

logger = logging.getLogger(__name__)

# Number of paragraphs before/after image run to collect as context
CONTEXT_PARA_RADIUS: int = 4


class DocxParser(BaseDocumentParser):
    """
    DOCX parser using python-docx.

    Two modes depending on available tooling:

    **Layout mode** (LibreOffice installed):
      Renders full DOCX pages to images at target DPI, preserving layout.
      surrouding_text is extracted from the same section as detected figures.

    **Embedded-image mode** (default, no LibreOffice needed):
      Extracts each embedded InlineShape image as a standalone ParsedPage.
      Surrounding paragraphs are collected from the paragraph-level context.

    Example:
        parser = DocxParser()
        doc = parser.parse(Path("report.docx"), dpi=150)
    """

    def parse(self, path: Path, dpi: int = 150) -> ParsedDocument:
        """
        Parse a DOCX file.

        Tries LibreOffice full-page rendering first; falls back to
        embedded-image extraction if LibreOffice is not available.

        Args:
            path: Absolute path to the .docx file.
            dpi:  Target resolution for LibreOffice rendering.

        Returns:
            ParsedDocument with one ParsedPage per extracted image.
        """
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            ) from exc

        doc = Document(str(path))
        document_title = self._extract_title(doc, path)

        logger.info(f"Parsing DOCX | path={path.name}")

        # Attempt full-page LibreOffice rendering
        if self._libreoffice_available():
            pages = self._parse_via_libreoffice(path, doc, dpi, document_title)
        else:
            pages = self._parse_embedded_images(doc, document_title)

        if not pages:
            # No images found: return whole-document text as single page
            pages = self._fallback_text_only_page(doc, document_title)

        return ParsedDocument(
            source_path=path,
            document_title=document_title,
            pages=pages,
        )

    # ------------------------------------------------------------------
    # LibreOffice full-page rendering
    # ------------------------------------------------------------------

    def _libreoffice_available(self) -> bool:
        """Check if LibreOffice (soffice) is on PATH."""
        try:
            result = subprocess.run(
                ["soffice", "--version"],
                capture_output=True,
                timeout=5,
            )
            return result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            return False

    def _parse_via_libreoffice(
        self,
        path: Path,
        doc: "Document",  # type: ignore[name-defined]
        dpi: int,
        document_title: Optional[str],
    ) -> List[ParsedPage]:
        """
        Render DOCX to PNG pages via LibreOffice headless mode.

        Args:
            path:           Source DOCX path.
            doc:            Opened python-docx Document.
            dpi:            Target render resolution.
            document_title: Pre-extracted document title.

        Returns:
            List of ParsedPage objects, one per page.
        """
        pages: List[ParsedPage] = []

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            cmd = [
                "soffice",
                "--headless",
                "--convert-to", "png",
                "--outdir", str(tmp_path),
                str(path),
            ]
            try:
                subprocess.run(cmd, capture_output=True, timeout=120, check=True)
            except subprocess.CalledProcessError as exc:
                logger.warning(
                    f"LibreOffice conversion failed | path={path.name} | error={exc}"
                )
                return []

            png_files = sorted(tmp_path.glob("*.png"))
            surrounding_texts = self._extract_page_texts(doc)

            for idx, png_file in enumerate(png_files, start=1):
                img_array = cv2.imread(str(png_file))
                if img_array is None:
                    continue

                ctx_text = surrounding_texts[idx - 1] if idx - 1 < len(surrounding_texts) else None
                pages.append(ParsedPage(
                    page_number=idx,
                    image_array=img_array,
                    source_format="docx",
                    is_scanned=False,
                    surrounding_text=ctx_text,
                    figure_caption=None,
                    document_title=document_title,
                ))

        return pages

    # ------------------------------------------------------------------
    # Embedded-image extraction (no LibreOffice)
    # ------------------------------------------------------------------

    def _parse_embedded_images(
        self,
        doc: "Document",  # type: ignore[name-defined]
        document_title: Optional[str],
    ) -> List[ParsedPage]:
        """
        Extract inline images from DOCX and collect surrounding paragraph text.

        Each InlineShape becomes one ParsedPage.

        Args:
            doc:            Opened python-docx Document.
            document_title: Pre-extracted document title.

        Returns:
            List of ParsedPage, one per embedded image.
        """
        from docx.oxml.ns import qn  # type: ignore[import-untyped]

        paragraphs = doc.paragraphs
        pages: List[ParsedPage] = []
        image_page_number = 0

        for para_idx, para in enumerate(paragraphs):
            # Check if paragraph contains any inline images
            image_rids = self._find_image_rids(para)
            if not image_rids:
                continue

            for rid in image_rids:
                image_bytes = self._get_image_bytes(doc, rid)
                if image_bytes is None:
                    continue

                img_array = self._bytes_to_array(image_bytes)
                if img_array is None:
                    continue

                image_page_number += 1

                # Collect context paragraphs
                surrounding, caption = self._collect_context(
                    paragraphs, para_idx
                )

                logger.debug(
                    f"DOCX embedded image | image_idx={image_page_number} | "
                    f"has_caption={caption is not None}"
                )

                pages.append(ParsedPage(
                    page_number=image_page_number,
                    image_array=img_array,
                    source_format="docx",
                    is_scanned=False,
                    surrounding_text=surrounding,
                    figure_caption=caption,
                    document_title=document_title,
                ))

        return pages

    def _find_image_rids(self, para: "Paragraph") -> List[str]:  # type: ignore[name-defined]
        """Return relationship IDs for all inline images in a paragraph."""
        from docx.oxml.ns import qn  # type: ignore[import-untyped]

        rids: List[str] = []
        for run in para.runs:
            # Inline images appear as <a:blip r:embed="rId..." />
            drawings = run._r.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            for blip in drawings:
                rid = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if rid:
                    rids.append(rid)
        return rids

    def _get_image_bytes(
        self,
        doc: "Document",  # type: ignore[name-defined]
        rid: str,
    ) -> Optional[bytes]:
        """Retrieve raw image bytes from the DOCX part by relationship ID."""
        try:
            part = doc.part.related_parts[rid]
            return part.blob
        except (KeyError, AttributeError):
            return None

    def _bytes_to_array(self, image_bytes: bytes) -> Optional[np.ndarray]:
        """Decode image bytes to BGR numpy array."""
        try:
            arr = np.frombuffer(image_bytes, dtype=np.uint8)
            img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
            return img
        except Exception:
            return None

    def _collect_context(
        self,
        paragraphs: list,
        image_para_idx: int,
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Collect surrounding paragraphs and identify any caption.

        Args:
            paragraphs:      Full paragraph list from the document.
            image_para_idx:  Index of the paragraph containing the image.

        Returns:
            (surrounding_text, figure_caption)
        """
        start = max(0, image_para_idx - CONTEXT_PARA_RADIUS)
        end = min(len(paragraphs), image_para_idx + CONTEXT_PARA_RADIUS + 1)

        texts: List[str] = []
        caption: Optional[str] = None

        for idx in range(start, end):
            if idx == image_para_idx:
                continue
            text = paragraphs[idx].text.strip()
            if not text:
                continue
            texts.append(text)
            # First caption-like paragraph wins
            if caption is None and self.is_caption(text):
                caption = text

        surrounding = self.truncate_context(self.clean_text("\n".join(texts))) if texts else None
        return surrounding, caption

    # ------------------------------------------------------------------
    # Fallback: no images found
    # ------------------------------------------------------------------

    def _fallback_text_only_page(
        self,
        doc: "Document",  # type: ignore[name-defined]
        document_title: Optional[str],
    ) -> List[ParsedPage]:
        """
        When no embedded images exist, return a single placeholder page
        with the document body text as surrounding_text.

        The image_array is a 1x1 white pixel (downstream stages will skip it
        as too small, so it effectively adds context metadata without images).
        """
        all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        surrounding = self.truncate_context(self.clean_text(all_text)) if all_text else None

        placeholder = np.full((100, 100, 3), 255, dtype=np.uint8)

        return [ParsedPage(
            page_number=1,
            image_array=placeholder,
            source_format="docx",
            is_scanned=False,
            surrounding_text=surrounding,
            figure_caption=None,
            document_title=document_title,
        )]

    # ------------------------------------------------------------------
    # Metadata helpers
    # ------------------------------------------------------------------

    def _extract_page_texts(self, doc: "Document") -> List[Optional[str]]:  # type: ignore[name-defined]
        """
        Extract per-page text blocks for LibreOffice mode.

        Since python-docx has no page-boundary concept, we approximate by
        splitting paragraphs evenly across pages.

        Returns:
            List of text strings (one per estimated page).
        """
        all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not all_texts:
            return []
        # Group into approx. 15-para pages
        page_size = 15
        groups: List[Optional[str]] = []
        for i in range(0, len(all_texts), page_size):
            chunk = all_texts[i : i + page_size]
            groups.append(self.truncate_context(self.clean_text("\n".join(chunk))))
        return groups

    def _extract_title(
        self,
        doc: "Document",  # type: ignore[name-defined]
        path: Path,
    ) -> Optional[str]:
        """
        Extract document title from DOCX core properties or first Heading 1.

        Args:
            doc:  Opened python-docx Document.
            path: Source file path (used for filename fallback).

        Returns:
            Title string or None.
        """
        # Core properties (File -> Properties -> Title)
        try:
            if doc.core_properties.title:
                return doc.core_properties.title.strip()
        except AttributeError:
            pass

        # First Heading 1 paragraph
        for para in doc.paragraphs:
            if para.style and "heading 1" in para.style.name.lower() and para.text.strip():
                return para.text.strip()

        # Fallback: filename
        return path.stem.replace("_", " ").replace("-", " ").title()
